#third-party imports
import boto3
from botocore.exceptions import NoCredentialsError
from docx.api import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
import matplotlib.pyplot as plt

#system imports
from io import BytesIO
import json
from datetime import datetime

#local/user imports


#Class Definitions
class DateTimeEncoder(json.JSONEncoder):
    def default(self, o):
        if isinstance(o, datetime):
            return o.isoformat()
        return super().default(o)
    

def add_heading(doc, text, style_name='Heading 1'):
    """
    Adds a heading to the document with the specified text and style.
    
    Args:
    - doc: Document object to which the heading will be added.
    - text: Text content of the heading.
    - style_name: Name of the style to be applied to the heading (default: 'Heading 1').
    """
    doc.add_heading(text, level=1).style = style_name

def add_subheading(doc, text, style_name='Heading 2'):
    """
    Adds a subheading to the document with the specified text and style.
    
    Args:
    - doc: Document object to which the subheading will be added.
    - text: Text content of the subheading.
    - style_name: Name of the style to be applied to the subheading (default: 'Heading 2').
    """
    doc.add_heading(text, level=2).style = style_name

def add_page_break(doc):
    """
    Adds a page break to the document.
    
    Args:
    - doc: Document object to which the page break will be added.
    """
    doc.add_page_break()

def add_empty_line(doc):
    """
    Adds an empty line (or new line) to the document.
    
    Args:
    - doc: Document object to which the empty line will be added.
    """
    doc.add_paragraph()

def add_paragraph(doc, text, style_name=None, alignment=None, bold=False, italic=False, underline=False):
    """
    Adds a paragraph of text to the document with specified formatting options.
    
    Args:
    - doc: Document object to which the paragraph will be added.
    - text: Text content of the paragraph.
    - style_name: Name of the style to be applied to the paragraph (default: None).
    - alignment: Alignment of the paragraph (default: None).
    - bold: Whether the text should be bold (default: False).
    - italic: Whether the text should be italicized (default: False).
    - underline: Whether the text should be underlined (default: False).
    """
    paragraph = doc.add_paragraph(text)
    
    # Apply style if specified
    if style_name:
        paragraph.style = style_name
    
    # Apply alignment if specified
    if alignment:
        paragraph.alignment = alignment
    
    # Apply formatting options
    run = paragraph.runs[0]
    run.bold = bold
    run.italic = italic
    run.underline = underline

# Example usage:

# # Create a new document
# doc = Document()

# # Add a heading
# add_heading(doc, 'Main Heading')

# # Add a subheading
# add_subheading(doc, 'Subheading')

# # Add a page break
# add_page_break(doc)

# # Add an empty line
# add_empty_line(doc)

# # Add a paragraph of text
# add_paragraph(doc, 'This is a paragraph of text.')

# # Save the document
# doc.save('example_document.docx')
