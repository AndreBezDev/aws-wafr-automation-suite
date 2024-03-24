#third-party imports
import boto3
from botocore.exceptions import NoCredentialsError
from docx.api import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches

import matplotlib.pyplot as plt
import numpy as np


#system imports
from io import BytesIO
import json
from datetime import datetime

#local/user imports
from docx_helper_functions import add_empty_line, add_heading, add_page_break, add_paragraph, add_subheading
from gpt_magic import quick_wins_section_prep

#Class Definitions
class DateTimeEncoder(json.JSONEncoder):
    def default(self, o):
        if isinstance(o, datetime):
            return o.isoformat()
        return super().default(o)


###Workflow
# Copy template to customer folder
# then edit the customer one
# Pull in AWS WAFR Report - extract data

# Function to create customer copy from base template
def copy_template_to_output_bucket(input_bucket, input_key, output_bucket, customer_folder):
    #To Do - create distinct files everytime

    #initialise S3 client
    s3 = boto3.client('s3')

    copy_source = {'Bucket': input_bucket, 'Key': input_key}
    output_key = f'{customer_folder}/CCL_WAFR_Report.docx'  # Adjust the output key as needed
    s3.copy_object(CopySource=copy_source, Bucket=output_bucket, Key=output_key)

    print(f"Template copied successfully to {output_bucket}/{output_key}")

# Function to download wip doc into container runtime
def download_docx_from_s3(bucket_name, key, local_path):
    #initialise S3 client
    s3 = boto3.client('s3')

    # Download the document from S3
    s3.download_file(bucket_name, key, local_path)

    #stored in container filesystem - no external volumes / mnt
    print(f"Document downloaded from S3 to {local_path}")


#store file in s3
def upload_to_s3(file_path, bucket_name, object_name):
    """
    Upload a file to an S3 bucket

    :param file_path: Path to the file to upload
    :param bucket_name: S3 bucket name
    :param object_name: S3 object name (key)
    :return: True if the file was uploaded successfully, else False
    """

    # Create an S3 client
    s3_client = boto3.client('s3')

    try:
        # Upload the file to S3
        s3_client.upload_file(file_path, bucket_name, object_name)
    except FileNotFoundError:
        print(f"The file '{file_path}' was not found.")
        return False
    except NoCredentialsError:
        print("AWS credentials not available.")
        return False
    
    print(f"File uploaded successfully to S3 bucket '{bucket_name}' with key '{object_name}'")
    return True

# Function to create a Word document with a heading and a table
def modify_word_document(local_path):
        
    # Create a new Word document
    doc = Document(f'{local_path}')

    # Add a heading "WAFR Risk Breakdown"
    doc.add_heading('WAFR Risk Breakdown Yeehoooowsa', level=1)

    # Save the document
    doc.save(f'{local_path}')
    print(f"Document modified at {local_path}")

#JSON to Word Table  
def json_to_table(json_data, doc, doc_path):
    """
    Convert JSON data to a table in a Word document with vertical cell merging and separated best practice choices.

    :param json_data: JSON data to convert to a table.
    :param doc: Word document object to write the table to.
    :param doc_path: Path to save the Word document.
    """
    # Initialize the table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    #pillar mapping
    pillar_mapping = {
        "operationalExcellence": "Operational Excellence",
        "security": "Security",
        "reliability": "Reliability",
        "performance": "Performance Efficiency",
        "costOptimization": "Cost Optimisation",
        "sustainability": "Sustainability"
    }

    # Add column headers and make them bold
    header_cells = table.rows[0].cells
    for idx, header_text in enumerate(['Pillar Name', 'Question Title', 'Best Practice Choice']):
        cell = header_cells[idx]
        cell.text = header_text
        # Apply bold formatting to the text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    merge_ranges = {}

    for pillar, questions in json_data.items():
        for question, details in questions.items():
            # Find the start and end row index for the current cell
            start_row_index = len(table.rows)
            for choice in details.get("UnselectedChoices", []):
                if choice != "None of these":
                    cells = table.add_row().cells
                    cells[0].text = pillar
                    cells[1].text = question
                    cells[2].text = choice.strip()
            
            end_row_index = len(table.rows)
            
            # Update merge ranges dictionary
            key = (pillar, question)
            if key in merge_ranges:
                merge_ranges[key] = (min(merge_ranges[key][0], start_row_index), max(merge_ranges[key][1], end_row_index))
            else:
                merge_ranges[key] = (start_row_index, end_row_index)

    # Merge cells vertically for each unique value
    for key, (start_row, end_row) in merge_ranges.items():
    # Replace the incorrect pillar name with the correct one
        pillar = key[0]
        correct_pillar = pillar_mapping.get(pillar, pillar)
        question = key[1]
        for col in range(2):  # Merge only first two columns
            start_cell = table.cell(start_row, col)
            end_cell = table.cell(end_row - 1, col)  # Subtract 1 to get the last cell
            start_text = start_cell.text
            # Replace the pillar name with the correct one
            start_text = start_text.replace(pillar, correct_pillar)
            start_cell.merge(end_cell)
            merged_cell = table.cell(start_row, col)  # Use start row for the merged cell
            merged_cell.text = start_text  # Set the text of the merged cell to the start cell's text

    # Set vertical alignment for all cells
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Save the document
    doc.save(doc_path)

# Create Risk Metrics breakdown / counts
def create_risk_metrics_table(json_data, doc_path, doc):
    """
    Creates a Word table for risk metrics breakdown.
    
    Args:
    - data: List of dictionaries containing risk metrics data.
    
    Returns:
    - Document: Word document object containing the generated table.
    """
    
    # Add a table with headers
    table = doc.add_table(rows=1, cols=len(json_data[0]) -1) #ignoring the None column
    table.style = 'Table Grid'

     # Add column headers and make them bold
    headers = list(json_data[0].keys())
    headers.remove('None')  # Remove 'None' column
    row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = row.cells[idx]
        if header == 'NotApplicable':
            cell.text = 'Not Applicable'
        else:
            cell.text = header
        # Apply bold formatting to the text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data to the table
    for entry in json_data:
        row = table.add_row().cells
        for idx, header in enumerate(headers):
            row[idx].text = str(entry[header])
    
    # Set alignment for all cells
    for row in table.rows:
        for cell in row.cells:
            # Set horizontal alignment to center
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Set vertical alignment to middle
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Set the alignment for the first column to left-aligned
    for cell in table.columns[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the document
    doc.save(doc_path)

def create_bar_chart(json_data, doc_path, doc):
    """
    Creates a bar chart of the risk metrics breakdown and inserts it into the Word document.
    
    Args:
    - json_data: List of dictionaries containing risk metrics data.
    - doc_path: Path to save the Word document.
    - doc: Word document object.
    """
    # Extract pillar names
    pillars = [entry['Name'] for entry in json_data]
    
    # Extract counts of HRIs and MRIs
    hris = [entry['High'] for entry in json_data]
    mris = [entry['Medium'] for entry in json_data]
    
    # Plot the bar chart
    x = np.arange(len(pillars))  # the label locations
    width = 0.35  # the width of the bars

    fig, ax = plt.subplots()
    rects1 = ax.bar(x - width/2, hris, width, label='HRIs', color='red')
    rects2 = ax.bar(x + width/2, mris, width, label='MRIs', color='orange')

    # Add some text for labels, title and custom x-axis tick labels, etc.
    ax.set_xlabel('Pillars')
    ax.set_ylabel('Counts')
    ax.set_title('Risk Metrics Breakdown')
    ax.set_xticks(x)
    ax.set_xticklabels(pillars, rotation=45, ha='right')
    ax.legend()

    # Attach a text label above each bar in rects
    def autolabel(rects):
        for rect in rects:
            height = rect.get_height()
            ax.annotate('{}'.format(height),
                        xy=(rect.get_x() + rect.get_width() / 2, height),
                        xytext=(0, 3),  # 3 points vertical offset
                        textcoords="offset points",
                        ha='center', va='bottom')

    autolabel(rects1)
    autolabel(rects2)

    # Adjust y-axis limit to create more vertical gap
    max_count = max(max(hris), max(mris)) + 1
    ax.set_ylim(0, max_count)

    fig.tight_layout()

    # Save the plot as an image
    plot_path = 'tmp/bar_chart.png'
    plt.savefig(plot_path)

    # Close the plot to release resources
    plt.close(fig)

    # Insert the saved image into the Word document
    doc.add_picture(plot_path, width=Inches(6))

    # Save the document
    doc.save(doc_path)

def create_qw_remediation_item(json_data, doc_path, doc):
    """
    Create a Word document section from a singular remediation plan part.

    Args:
    - json_data: JSON object containing the data to be written to the document.
    - output_path: Path to save the Word document.
    """

    #add subheading i.e. 1. Implement application telemetry
    text = f"{json_data['best_practice_option']}"
    add_subheading(doc, text, style_name='Heading 4')
    
    #add sub-subheading 3 i.e. Description
    text = "Description"
    add_subheading(doc, text, style_name='Heading 5')

    text = f"{json_data['remediation_description']}"
    add_paragraph(doc, text, style_name=None, alignment=None, bold=False, italic=False, underline=False)

    # Add spacing between sections
    doc.add_paragraph()

    #add sub-subheading 3 i.e. Solution
    text = "Remediation"
    add_subheading(doc, text, style_name='Heading 5')

    text = f"{json_data['remediation_solution']}"
    add_paragraph(doc, text, style_name=None, alignment=None, bold=False, italic=False, underline=False)

    # Add spacing between sections
    doc.add_paragraph()

    #add sub-subheading 3 i.e. General Considerations
    text = "Additional Considerations"
    add_subheading(doc, text, style_name='Heading 5')

    text = f"{json_data['remediation_general_considerations']}"
    add_paragraph(doc, text, style_name=None, alignment=None, bold=False, italic=False, underline=False)

    # Add spacing between sections
    doc.add_paragraph()

    #add sub-subheading 3 i.e. Effort Estimate
    text = "Effort Estimate"
    add_subheading(doc, text, style_name='Heading 5')

    text = f"{json_data['effort_estimate']}"
    add_paragraph(doc, text, style_name=None, alignment=None, bold=False, italic=False, underline=False)

    # Add spacing between sections
    doc.add_paragraph()

    #add sub-subheading 3 i.e. Resourcing and Skills
    text = "Resourcing and Skills"
    add_subheading(doc, text, style_name='Heading 5')

    text = f"{json_data['resources_needed']}"
    add_paragraph(doc, text, style_name=None, alignment=None, bold=False, italic=False, underline=False)

    # Add spacing between sections
    doc.add_paragraph()

    #add sub-subheading 3 i.e. Domain Impact
    text = "Domain Impact"
    add_subheading(doc, text, style_name='Heading 5')

    text = f"{json_data['domain_impact']}"
    add_paragraph(doc, text, style_name=None, alignment=None, bold=False, italic=False, underline=False)

    # Add spacing between sections
    doc.add_paragraph()
    # Save the document
    doc.save(doc_path)


def create_top_10_qw_table(json_data, doc, doc_path):
    # Add a table with headers
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Add column headers
    headers = ["#", "Question", "Best Practice"]
    row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = row.cells[idx]
        cell.text = header
        # Apply bold formatting to the text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data to the table
    for entry in json_data:
        row = table.add_row().cells
        row[0].text = str(entry["quick_win_id"])
        row[1].text = entry["best_practice_question"]
        row[2].text = entry["unselected_best_practice_item"]
    
    # Set alignment for all cells
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Horizontal alignment
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Vertical alignment

    # Adjust width of the first column
    table.autofit = True
    table.columns[0].width = Inches(0.5)  # Set width to fit 2digits
    table.columns[1].width = Inches(3.0)  
    table.columns[2].width = Inches(3.5)  

    # Save the document
    doc.save(doc_path)
    

def initialise_docs(inputBucket, inputKey, outputBucket,customerFoler):
    #initialise clients
    s3 = boto3.client('s3')

     #config vars
    input_bucket = inputBucket
    input_key = inputKey
    output_bucket = outputBucket  # Replace with your output bucket name
    customer_folder = customerFoler  # Specify customer-specific folder name
    
    # Copy the template to the output bucket under the customer-specific folder
    copy_template_to_output_bucket(input_bucket, input_key, output_bucket, customer_folder)
    
    # Specify S3 bucket details and local file path
    bucket_name = f'{output_bucket}'  # Replace with your S3 bucket name
    key = f'{customer_folder}/CCL_WAFR_Report.docx'      # Replace with the key of the document in your S3 bucket
    local_path = 'tmp/wip_document.docx' # Specify the local file path where the document will be saved

    # Download the customer document copy from S3
    download_docx_from_s3(bucket_name, key, local_path)

    # Open Doc for WIP
    doc_path = f'{local_path}'
    doc = Document(doc_path)
    
    #insert Risk Metrics Breakdown
    add_heading(doc,'Risk Breakdown by Pillar', style_name='Heading 1') #insert Risk Metrics Header
    add_empty_line(doc) #insert empty line
    json_file_path = 'tmp/json/risk_metrics.json'
    
    with open(json_file_path) as json_file:
        json_data = json.load(json_file)
    
    create_risk_metrics_table(json_data, doc_path, doc)
    add_empty_line(doc) #insert empty line

    #insert bar chart
    create_bar_chart(json_data, doc_path, doc)

    # insert HRI table
    add_page_break(doc)
    add_heading(doc,'High Risk Items', style_name='Heading 1') #insert HRI Header
    add_empty_line(doc) #insert empty line
    json_file_path = 'tmp/json/filter_high_risk_questions.json'
    
    with open(json_file_path) as json_file:
        json_data = json.load(json_file)

    json_to_table(json_data, doc, doc_path)
    add_empty_line(doc) #insert empty line

    #insert MRI Table in Doc
    add_page_break(doc)
    add_heading(doc,'Medium Risk Items', style_name='Heading 1') #insert MRI Header
    add_empty_line(doc) #insert empty line
    json_file_path = 'tmp/json/filter_medium_risk_questions.json'
    
    with open(json_file_path) as json_file:
        json_data = json.load(json_file)
    
    json_to_table(json_data, doc, doc_path)
    add_empty_line(doc) #insert empty line

    # Start Remediation Plan
    add_page_break(doc)
    add_heading(doc,'Remediation Plan', style_name='Heading 2')
    #Insert Quick Wins Section
    add_heading(doc,'Quick Wins', style_name='Heading 3')
    remediation_pts = quick_wins_section_prep()

    # Insert QW Table
    json_file_path = 'tmp/json/top_10_quick_wins.json'
    
    with open(json_file_path) as json_file:
        json_data = json.load(json_file)
    create_top_10_qw_table(json_data, doc, doc_path)
    add_empty_line(doc) #insert empty line

    # Insert QW Remediation sections
    for file in remediation_pts:
        print(f"current file path is: {file}")
        with open(file) as json_file:
            json_data = json.load(json_file)
        create_qw_remediation_item(json_data, doc_path, doc)  
        add_empty_line(doc)

    #store WIP doc in S3
    upload_to_s3(local_path, bucket_name, key)


    # Modify the Word document
    #modify_word_document(local_path)

    # # Generate the bar chart
    # generate_bar_chart()

    # print("Word document and bar chart generated successfully.")

def test():
    # Define workload params - To Do: pull from incoming Request POST 
    workload_id = '66658f2b4909462a1bc9a50349bcaec7' #To Do - fetch appropriate id from API
    lens_alias = 'arn:aws:wellarchitected::aws:lens/wellarchitected'
    milestone_number = 1 #normally 1, check if not 2 - want to get latest

    # Define s3 artefact params - To Do enable selection of templates 
    input_bucket = 'aws-wafr-automation-base-templates'
    input_key = 'CCL-2024/CCL AWS WAFR - Report Template v0.1.docx'
    output_bucket = 'aws-wafr-automation-output-reports'  # Replace with your output bucket name
    customer_folder = 'KMD'  # Specify customer-specific folder name
    
    #initialise clients
    s3 = boto3.client('s3')
    
    # Copy the template to the output bucket under the customer-specific folder
    #copy_template_to_output_bucket(input_bucket, input_key, output_bucket, customer_folder)
    
    # Specify S3 bucket details and local file path
    bucket_name = f'{output_bucket}'  # Replace with your S3 bucket name
    key = f'{customer_folder}/CCL_WAFR_Report.docx'      # Replace with the key of the document in your S3 bucket
    local_path = 'tmp/wip_document.docx' # Specify the local file path where the document will be saved

    # Download the customer document copy from S3
    #download_docx_from_s3(bucket_name, key, local_path)

    # Open Doc for WIP
    doc_path = f'{local_path}'
    doc = Document(doc_path)

    add_page_break(doc)
    add_heading(doc,'Testing', style_name='Heading 2')

    json_file_path = 'tmp/json/top_10_quick_wins.json'
    
    with open(json_file_path) as json_file:
        json_data = json.load(json_file)
    create_top_10_qw_table(json_data, doc, doc_path)

    # Save the document
    doc.save(doc_path)

#if __name__ == "__main__":
    #test()